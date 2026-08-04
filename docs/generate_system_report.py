"""Generate the illustrated Waste Picker System report and its source graphs.

The report is built from the dated, authenticated API snapshot in
``report_data_snapshot.json`` and the verified dashboard/mobile evidence in
``docs/screenshots``. It deliberately describes only implemented features.
"""

from __future__ import annotations

import json
import math
from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import Ellipse, FancyArrowPatch, FancyBboxPatch, Rectangle
from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SNAPSHOT_PATH = ROOT / "report_data_snapshot.json"
SCREENSHOTS = ROOT / "screenshots"
GENERATED = ROOT / "generated"
OUTPUT = ROOT / "Waste Picker System - Detailed System Report.docx"

GREEN = "1B5E20"
GREEN_LIGHT = "E8F5E9"
BLUE = "1565C0"
BLUE_LIGHT = "E3F2FD"
AMBER = "D97706"
AMBER_LIGHT = "FFF7ED"
RED = "C62828"
INK = "17211A"
MUTED = "5B6B60"
LINE = "DDE6DF"
WHITE = "FFFFFF"

MPL_GREEN = "#2E7D32"
MPL_BLUE = "#1976D2"
MPL_AMBER = "#F0B429"
MPL_RED = "#C62828"
MPL_MUTED = "#718078"


def load_snapshot() -> dict:
    with SNAPSHOT_PATH.open("r", encoding="utf-8-sig") as handle:
        data = json.load(handle)
    required = {
        "captured_at",
        "total_pickers",
        "approved",
        "pending",
        "rejected",
        "suspended",
        "total_kg",
        "by_region",
        "by_material",
        "registration_trend",
    }
    missing = sorted(required - data.keys())
    if missing:
        raise ValueError(f"Snapshot is missing: {', '.join(missing)}")
    return data


def as_of(data: dict) -> str:
    stamp = datetime.fromisoformat(data["captured_at"])
    return stamp.strftime("%d %B %Y at %H:%M EAT")


def save_figure(fig: plt.Figure, name: str) -> Path:
    path = GENERATED / name
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def chart_style(ax: plt.Axes, title: str, subtitle: str = "") -> None:
    ax.set_title(title, loc="left", fontsize=18, weight="bold", color="#153D1B", pad=22)
    if subtitle:
        ax.text(0, 1.02, subtitle, transform=ax.transAxes, fontsize=10.5, color="#5B6B60")
    ax.spines[["top", "right"]].set_visible(False)
    ax.spines[["left", "bottom"]].set_color("#DDE6DF")
    ax.tick_params(colors="#526158", labelsize=9)
    ax.grid(axis="y", color="#E8EEE9", linewidth=0.8, zorder=0)


def generate_charts(data: dict) -> None:
    statuses = ["Approved", "Pending", "Rejected", "Suspended"]
    values = [data["approved"], data["pending"], data["rejected"], data["suspended"]]
    colors = [MPL_GREEN, MPL_AMBER, MPL_RED, MPL_MUTED]
    fig, ax = plt.subplots(figsize=(10, 5.5))
    wedges, _ = ax.pie(
        values,
        colors=colors,
        startangle=90,
        counterclock=False,
        wedgeprops={"width": 0.34, "edgecolor": "white", "linewidth": 3},
    )
    ax.text(0, 0.08, f"{data['total_pickers']}", ha="center", va="center", fontsize=30, weight="bold", color="#153D1B")
    ax.text(0, -0.16, "registered", ha="center", va="center", fontsize=11, color="#5B6B60")
    legend = [f"{name}: {value} ({(value / data['total_pickers'] * 100 if data['total_pickers'] else 0):.1f}%)" for name, value in zip(statuses, values)]
    ax.legend(wedges, legend, loc="center left", bbox_to_anchor=(0.98, 0.5), frameon=False, fontsize=11)
    ax.set_title("Registration status summary", fontsize=18, weight="bold", color="#153D1B", pad=18)
    ax.text(0.5, -0.02, f"Authenticated snapshot · {as_of(data)}", transform=fig.transFigure, ha="center", fontsize=9, color="#718078")
    save_figure(fig, "status_summary.png")

    regions = data["by_region"]
    labels = [row["region"] for row in regions]
    approved = [row["approved"] for row in regions]
    pending = [row["pending"] for row in regions]
    rejected = [row["rejected"] for row in regions]
    suspended = [row["suspended"] for row in regions]
    x = list(range(len(labels)))
    fig, ax = plt.subplots(figsize=(12, 6.2))
    ax.bar(x, approved, color=MPL_GREEN, label="Approved", zorder=3)
    ax.bar(x, pending, bottom=approved, color=MPL_AMBER, label="Pending", zorder=3)
    bottom = [a + p for a, p in zip(approved, pending)]
    ax.bar(x, rejected, bottom=bottom, color=MPL_RED, label="Rejected", zorder=3)
    bottom2 = [a + p + r for a, p, r in zip(approved, pending, rejected)]
    ax.bar(x, suspended, bottom=bottom2, color=MPL_MUTED, label="Suspended", zorder=3)
    chart_style(ax, "Registrations by region", "Status distribution across all ten configured counties")
    ax.set_xticks(x, labels, rotation=28, ha="right")
    ax.set_ylabel("Waste picker records")
    ax.yaxis.set_major_locator(plt.MaxNLocator(integer=True))
    ax.legend(ncol=4, frameon=False, loc="upper right")
    fig.tight_layout()
    save_figure(fig, "registrations_by_region.png")

    materials = data["by_material"]
    labels = [row["material"] for row in materials]
    weights = [float(row["kg"]) for row in materials]
    palette = ["#00897B", "#1976D2", "#EF6C00", "#795548", "#7B1FA2", "#2E7D32"]
    fig, ax = plt.subplots(figsize=(10.5, 6))
    wedges, _ = ax.pie(
        weights,
        startangle=90,
        counterclock=False,
        colors=palette,
        wedgeprops={"width": 0.38, "edgecolor": "white", "linewidth": 2.5},
    )
    ax.text(0, 0.07, f"{data['total_kg']:,.2f}", ha="center", va="center", fontsize=23, weight="bold", color="#153D1B")
    ax.text(0, -0.15, "kilograms", ha="center", va="center", fontsize=10.5, color="#5B6B60")
    legend = [f"{label}: {weight:,.2f} kg ({weight / data['total_kg'] * 100:.1f}%)" for label, weight in zip(labels, weights)]
    ax.legend(wedges, legend, loc="center left", bbox_to_anchor=(0.96, 0.5), frameon=False, fontsize=10.5)
    ax.set_title("Recorded waste by material", fontsize=18, weight="bold", color="#153D1B", pad=18)
    save_figure(fig, "material_mix.png")

    trend = data["registration_trend"]
    months = [datetime.strptime(row["month"], "%Y-%m").strftime("%b %Y") for row in trend]
    counts = [row["registrations"] for row in trend]
    fig, ax = plt.subplots(figsize=(11.5, 5.5))
    ax.plot(months, counts, color=MPL_GREEN, linewidth=3, marker="o", markersize=8, zorder=3)
    ax.fill_between(months, counts, color="#C8E6C9", alpha=0.65, zorder=2)
    for index, value in enumerate(counts):
        ax.annotate(str(value), (index, value), xytext=(0, 10), textcoords="offset points", ha="center", weight="bold", color="#1B5E20")
    chart_style(ax, "Six-month registration trend", "New waste picker records created per calendar month")
    ax.set_ylabel("New registrations")
    ax.yaxis.set_major_locator(plt.MaxNLocator(integer=True))
    ax.set_ylim(bottom=0)
    fig.tight_layout()
    save_figure(fig, "registration_trend.png")


def box(ax: plt.Axes, x: float, y: float, w: float, h: float, text: str, *, fc="#E8F5E9", ec="#2E7D32", size=10, radius=0.02) -> None:
    patch = FancyBboxPatch((x, y), w, h, boxstyle=f"round,pad=0.012,rounding_size={radius}", facecolor=fc, edgecolor=ec, linewidth=1.8)
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=size, color="#173B20", weight="bold", wrap=True)


def arrow(ax: plt.Axes, start: tuple[float, float], end: tuple[float, float], *, color="#477052", text: str | None = None, offset=(0, 0)) -> None:
    patch = FancyArrowPatch(start, end, arrowstyle="-|>", mutation_scale=13, linewidth=1.5, color=color)
    ax.add_patch(patch)
    if text:
        ax.text((start[0] + end[0]) / 2 + offset[0], (start[1] + end[1]) / 2 + offset[1], text, fontsize=8, color="#526158", ha="center", bbox={"facecolor": "white", "edgecolor": "none", "pad": 1})


def diagram_canvas(title: str, figsize=(12, 7)) -> tuple[plt.Figure, plt.Axes]:
    fig, ax = plt.subplots(figsize=figsize)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.965, title, ha="center", va="top", fontsize=19, weight="bold", color="#174A26")
    return fig, ax


def generate_diagrams() -> None:
    fig, ax = diagram_canvas("Use case diagram — actors and implemented capabilities", (12, 8))
    ax.add_patch(Rectangle((0.25, 0.08), 0.5, 0.78, facecolor="#F7FBF8", edgecolor=MPL_GREEN, linewidth=2))
    ax.text(0.5, 0.825, "Waste Picker Management System", ha="center", fontsize=13, weight="bold", color="#1B5E20")
    picker_cases = [(0.34, 0.66, "Register / sign in"), (0.34, 0.52, "View / update profile"), (0.34, 0.38, "View digital identity"), (0.34, 0.24, "Receive messages"), (0.34, 0.10, "Record collections")]
    admin_cases = [(0.56, 0.66, "Approve / suspend records"), (0.56, 0.52, "Assign roles"), (0.56, 0.38, "Broadcast messages"), (0.56, 0.24, "Monitor activity"), (0.56, 0.10, "View / export reports")]
    box(ax, 0.015, 0.46, 0.15, 0.12, "WASTE\nPICKER", fc="#E3F2FD", ec=MPL_BLUE, size=10.5)
    box(ax, 0.835, 0.46, 0.15, 0.12, "ADMINISTRATOR", fc="#E8F5E9", ec=MPL_GREEN, size=9.7)
    for x, y, label in picker_cases:
        ellipse = Ellipse((x + 0.06, y + 0.045), 0.2, 0.09, facecolor="#E3F2FD", edgecolor=MPL_BLUE, linewidth=1.5)
        ax.add_patch(ellipse)
        ax.text(x + 0.06, y + 0.045, label, ha="center", va="center", fontsize=8.6, weight="bold", color="#174A74")
        ax.plot([0.165, x - 0.04], [0.52, y + 0.045], color=MPL_BLUE, linewidth=1)
    for x, y, label in admin_cases:
        ellipse = Ellipse((x + 0.05, y + 0.045), 0.2, 0.09, facecolor="#E8F5E9", edgecolor=MPL_GREEN, linewidth=1.5)
        ax.add_patch(ellipse)
        ax.text(x + 0.05, y + 0.045, label, ha="center", va="center", fontsize=8.6, weight="bold", color="#20572B")
        ax.plot([x + 0.15, 0.835], [y + 0.045, 0.52], color=MPL_GREEN, linewidth=1)
    save_figure(fig, "use_case.png")

    fig, ax = diagram_canvas("Four-tier system architecture", (12, 8))
    layers = [
        (0.76, "PRESENTATION", "#E3F2FD", "#1976D2", ["Android mobile app\n(Kotlin + XML)", "React admin dashboard\n(Vite + Recharts)"]),
        (0.56, "APPLICATION", "#EDF4FB", "#25557E", ["Express REST API", "Authentication + role guards", "ID assignment workflow", "Notification fan-out"]),
        (0.36, "INTEGRATION", "#E8F5E9", "#258153", ["Supabase Auth", "Supabase Storage", "HTTPS / JSON"]),
        (0.16, "DATA", "#FFF3E8", "#C55C0A", ["PostgreSQL tables", "Reporting views", "Triggers + RLS", "picker-photos bucket"]),
    ]
    for y, title, fc, ec, items in layers:
        ax.add_patch(Rectangle((0.05, y - 0.1), 0.9, 0.17, facecolor=fc, edgecolor=ec, linewidth=2))
        ax.text(0.07, y + 0.045, title, fontsize=11, weight="bold", color=ec)
        item_w = 0.18 if len(items) == 4 else 0.34
        gap = (0.82 - item_w * len(items)) / max(1, len(items) - 1)
        for i, item in enumerate(items):
            x = 0.09 + i * (item_w + gap)
            box(ax, x, y - 0.065, item_w, 0.09, item, fc="white", ec=ec, size=8.7, radius=0.008)
    for y1, y2 in [(0.66, 0.62), (0.46, 0.42), (0.26, 0.22)]:
        arrow(ax, (0.5, y1), (0.5, y2), color="#526158")
        arrow(ax, (0.53, y2), (0.53, y1), color="#526158")
    save_figure(fig, "architecture.png")

    fig, ax = diagram_canvas("Data flow — operational information movement", (12, 7.5))
    box(ax, 0.03, 0.62, 0.14, 0.12, "Waste picker", fc="#E3F2FD", ec=MPL_BLUE)
    box(ax, 0.83, 0.62, 0.14, 0.12, "Administrator", fc="#E8F5E9", ec=MPL_GREEN)
    box(ax, 0.25, 0.66, 0.2, 0.11, "Registration & profile\nmanagement", fc="#F3F8FC", ec="#25557E")
    box(ax, 0.55, 0.66, 0.2, 0.11, "Approval & role\nmanagement", fc="#F3F8FC", ec="#25557E")
    box(ax, 0.25, 0.39, 0.2, 0.11, "Collection activity", fc="#F3F8FC", ec="#25557E")
    box(ax, 0.55, 0.39, 0.2, 0.11, "Communication fan-out", fc="#F3F8FC", ec="#25557E")
    box(ax, 0.4, 0.16, 0.2, 0.11, "Reporting & analytics", fc="#FFF7ED", ec=MPL_AMBER)
    box(ax, 0.4, 0.53, 0.2, 0.09, "PostgreSQL + Auth + Storage", fc="#E8F5E9", ec=MPL_GREEN)
    arrow(ax, (0.17, 0.68), (0.25, 0.71), text="registration/profile")
    arrow(ax, (0.75, 0.71), (0.83, 0.68), text="review/actions")
    arrow(ax, (0.45, 0.71), (0.55, 0.71), text="pending record")
    arrow(ax, (0.35, 0.66), (0.45, 0.62))
    arrow(ax, (0.65, 0.66), (0.55, 0.62))
    arrow(ax, (0.17, 0.64), (0.25, 0.45), text="material + kg", offset=(0, -0.02))
    arrow(ax, (0.75, 0.45), (0.83, 0.64), text="delivery/read totals", offset=(0, -0.02))
    arrow(ax, (0.83, 0.62), (0.75, 0.45), text="broadcast")
    arrow(ax, (0.55, 0.43), (0.17, 0.62), text="inbox message", offset=(0, 0.02))
    arrow(ax, (0.35, 0.39), (0.45, 0.27))
    arrow(ax, (0.65, 0.39), (0.55, 0.27))
    arrow(ax, (0.5, 0.53), (0.5, 0.27), color=MPL_GREEN, text="tables + views", offset=(0.06, 0))
    arrow(ax, (0.6, 0.21), (0.83, 0.62), text="summaries")
    save_figure(fig, "data_flow.png")

    fig, ax = diagram_canvas("Registration-to-identity workflow", (8.5, 12))
    steps = [
        (0.82, "Open app and submit registration", "#E3F2FD", MPL_BLUE),
        (0.69, "API validates input and creates Auth identity", "#EDF4FB", "#25557E"),
        (0.56, "Profile stored with status = pending", "#FFF7ED", MPL_AMBER),
        (0.43, "Administrator reviews the registration", "#E8F5E9", MPL_GREEN),
        (0.30, "Approval triggers atomic region/year sequence", "#E8F5E9", MPL_GREEN),
        (0.17, "WP-<REGION>-<YEAR>-<NNNN> assigned", "#DFF4E4", MPL_GREEN),
        (0.04, "Profile refresh unlocks the digital identity card", "#DFF4E4", MPL_GREEN),
    ]
    for index, (y, label, fc, ec) in enumerate(steps):
        box(ax, 0.2, y, 0.6, 0.085, label, fc=fc, ec=ec, size=10.5)
        if index < len(steps) - 1:
            arrow(ax, (0.5, y), (0.5, steps[index + 1][0] + 0.085), color="#477052")
    ax.text(0.83, 0.45, "Rejected / suspended\n→ explanatory status screen", ha="left", va="center", fontsize=9, color=MPL_RED)
    arrow(ax, (0.8, 0.47), (0.8, 0.56), color=MPL_RED)
    save_figure(fig, "registration_flow.png")

    fig, ax = diagram_canvas("Entity relationship diagram — implemented data model", (13, 8))
    entities = {
        "regions": (0.03, 0.68, "regions\nPK id · name · code"),
        "waste_pickers": (0.29, 0.68, "waste_pickers\nPK/FK id · picker_id\nstatus · role · region_id"),
        "admins": (0.73, 0.68, "admins\nPK/FK id · email · role"),
        "collections": (0.03, 0.36, "collections\nPK id · FK picker_id\nmaterial · weight_kg"),
        "announcements": (0.55, 0.36, "announcements\nPK id · created_by\naudience · region_id"),
        "recipients": (0.29, 0.08, "announcement_recipients\nFK announcement_id · picker_id\nread_at"),
        "tokens": (0.03, 0.08, "device_tokens\nPK id · FK picker_id · token"),
        "sequences": (0.73, 0.08, "id_sequences\nPK region_code + year\nlast_value"),
    }
    for key, (x, y, label) in entities.items():
        width = 0.23 if key in {"waste_pickers", "announcements", "recipients"} else 0.2
        box(ax, x, y, width, 0.13, label, fc="#F7FBF8", ec=MPL_GREEN, size=8.6, radius=0.008)
    arrow(ax, (0.23, 0.745), (0.29, 0.745), text="1:N", offset=(0, 0.025))
    arrow(ax, (0.39, 0.68), (0.18, 0.49), text="1:N", offset=(0, 0.02))
    arrow(ax, (0.52, 0.72), (0.73, 0.72), text="approved by", offset=(0, 0.02))
    arrow(ax, (0.83, 0.68), (0.67, 0.49), text="creates", offset=(0, 0.02))
    arrow(ax, (0.55, 0.39), (0.47, 0.21), text="1:N", offset=(0, 0.02))
    arrow(ax, (0.4, 0.68), (0.4, 0.21), text="receives", offset=(0.045, 0))
    arrow(ax, (0.29, 0.72), (0.23, 0.15), text="has tokens", offset=(-0.035, 0))
    arrow(ax, (0.39, 0.68), (0.82, 0.21), text="ID trigger", offset=(0, 0.02))
    ax.text(0.82, 0.52, "Reporting views aggregate\nregistrations, communication\nand collection activity", ha="center", va="center", fontsize=9.5, color="#526158", bbox={"boxstyle": "round,pad=0.5", "facecolor": "#FFF7ED", "edgecolor": MPL_AMBER})
    save_figure(fig, "erd.png")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8.5)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def configure_document(doc: Document, data: dict) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for style_name, size, color in [("Title", 30, GREEN), ("Heading 1", 20, GREEN), ("Heading 2", 14, BLUE), ("Heading 3", 11.5, GREEN)]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "WASTE PICKER MANAGEMENT SYSTEM  ·  DETAILED SYSTEM REPORT"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(7.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    footer.add_run(f"Authenticated system snapshot: {as_of(data)}   ·   ")
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_number(footer)

    doc.core_properties.title = "Waste Picker Management System — Detailed System Report"
    doc.core_properties.subject = "INSY 492 Senior Project system documentation and operational summary"
    doc.core_properties.author = "Wiclife Omondi Ongo"
    doc.core_properties.keywords = "waste picker, management system, React, Express, Supabase, system report"
    doc.core_properties.comments = "Reproducibly generated from the repository evidence and dated API snapshot."


def add_paragraph(doc: Document, text: str, *, bold_prefix: str | None = None, align=None) -> None:
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    header_cells = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    for i, heading in enumerate(headers):
        set_cell_shading(header_cells[i], GREEN)
        header_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = header_cells[i].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(str(heading))
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        run.font.size = Pt(8.6)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 1:
                set_cell_shading(cells[i], "F7FAF8")
            paragraph = cells[i].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(str(value))
            run.font.size = Pt(8.4)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_image(doc: Document, path: Path, caption: str, *, width=6.65, page_break=False, max_height=8.35) -> None:
    if not path.is_file():
        raise FileNotFoundError(path)
    with Image.open(path) as source_image:
        aspect_ratio = source_image.width / source_image.height
    fitted_width = min(width, max_height * aspect_ratio)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run().add_picture(str(path), width=Inches(fitted_width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    if page_break:
        doc.add_page_break()


def add_mobile_pair(doc: Document, left: tuple[Path, str], right: tuple[Path, str]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for cell, (path, caption) in zip(table.rows[0].cells, [left, right]):
        if not path.is_file():
            raise FileNotFoundError(path)
        cell.width = Inches(3.2)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        picture_paragraph = cell.paragraphs[0]
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.add_run().add_picture(str(path), width=Inches(2.72))
        caption_paragraph = cell.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_paragraph.add_run(caption)
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def heading_page(doc: Document, title: str, subtitle: str | None = None) -> None:
    doc.add_page_break()
    doc.add_heading(title, level=1)
    if subtitle:
        paragraph = doc.add_paragraph(subtitle)
        paragraph.style = doc.styles["Subtitle"]


def build_report(data: dict) -> None:
    doc = Document()
    configure_document(doc, data)

    # Cover
    doc.add_paragraph().paragraph_format.space_after = Pt(65)
    eyebrow = doc.add_paragraph("UNIVERSITY OF EASTERN AFRICA, BARATON")
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in eyebrow.runs:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(BLUE)
    sub = doc.add_paragraph("School of Business · Department of Information Systems and Computing")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(32)
    title = doc.add_paragraph("DETAILED SYSTEM REPORT", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("A Web-Based Mobile Waste Picker App")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.bold = True
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(30)
    for line in [
        "Wiclife Omondi Ongo",
        "Bachelor of Business Information Technology · INSY 492 Senior Project",
        "Supervisor: Dr. Victor Mony",
        "Academic Year 2025/2026",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(11 if line == "Wiclife Omondi Ongo" else 9.5)
        p.runs[0].font.bold = line == "Wiclife Omondi Ongo"
    doc.add_paragraph().paragraph_format.space_after = Pt(32)
    live = doc.add_paragraph("Live system: https://waste-picker-system.vercel.app")
    live.alignment = WD_ALIGN_PARAGRAPH.CENTER
    live.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    snap = doc.add_paragraph(f"Operational data as of {as_of(data)}")
    snap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    snap.runs[0].font.size = Pt(9)
    snap.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    heading_page(doc, "Contents")
    add_paragraph(doc, "This generated contents list mirrors the report's numbered sections. Word heading styles are applied throughout for navigation and PDF bookmarks.")
    contents = [
        "0. Executive system summary",
        "1. System overview and requirements",
        "2. System architecture and data flow",
        "3. Database design and unique identity generation",
        "4. Authentication, authorisation and security model",
        "5. Administrative dashboard",
        "6. Mobile application",
        "7. Reporting and operational analytics",
        "8. Deployment and operations",
        "9. Verification and testing",
        "10. Security assessment",
        "11. Limitations and future work",
        "12. Conclusion",
        "Appendix A. API and repository reference",
        "Appendix B. Report provenance",
    ]
    add_numbered(doc, contents)

    heading_page(doc, "0. Executive system summary", "Current operational position, generated from authenticated system data")
    total = data["total_pickers"]
    approval_rate = data["approved"] / total * 100 if total else 0
    read_rate = data["messages_read"] / data["messages_delivered"] * 100 if data["messages_delivered"] else 0
    add_paragraph(doc, f"The Waste Picker Management System is deployed and operational. At {as_of(data)}, it held {total} waste picker records across ten configured counties. Thirteen records were approved ({approval_rate:.1f}%), three awaited review, and none were rejected or suspended. Approval assigns a permanent, region-aware identifier and unlocks the mobile identity card and activity workflow.")
    add_paragraph(doc, f"The activity module contained {data['collection_records']} collection records from {data['active_collectors']} active collectors, representing {data['total_kg']:,.2f} kg of material. The communication module contained {data['announcements']} broadcasts, {data['messages_delivered']} recipient deliveries and {data['messages_read']} recorded read, a {read_rate:.1f}% read rate. Communication engagement is therefore the clearest operational improvement opportunity identified by this snapshot.")
    add_table(
        doc,
        ["Indicator", "Value", "Management meaning"],
        [
            ["Registered waste pickers", total, "Coverage of the community register"],
            ["Approved", f"{data['approved']} ({approval_rate:.1f}%)", "Profiles with an issued identity"],
            ["Pending", data["pending"], "Administrator review queue"],
            ["Waste recorded", f"{data['total_kg']:,.2f} kg", f"{data['collection_records']} records; {data['active_collectors']} active collectors"],
            ["Broadcasts", data["announcements"], f"{data['messages_delivered']} deliveries; {data['messages_read']} reads"],
        ],
    )
    add_image(doc, GENERATED / "status_summary.png", "Figure 1. Registration status distribution at the report snapshot.", width=6.25)
    add_image(doc, GENERATED / "registrations_by_region.png", "Figure 2. Regional distribution, separated by registration status.", width=6.65, page_break=True)
    add_image(doc, GENERATED / "material_mix.png", "Figure 3. Recorded material mix by cumulative kilograms.", width=6.35)
    add_image(doc, GENERATED / "registration_trend.png", "Figure 4. New registrations during the six-month reporting window.", width=6.65)

    heading_page(doc, "1. System overview and requirements")
    add_paragraph(doc, "Waste pickers across Western Kenya contribute materially to recycling and cleaner communities but often remain outside formal registers. Without a trusted identity record, coordinators struggle to identify beneficiaries, communicate urgent information, measure collected material or demonstrate programme reach. The project addresses that gap with one shared platform exposed through a mobile application for waste pickers and a web dashboard for coordinators.")
    doc.add_heading("1.1 Objectives", level=2)
    add_bullets(doc, [
        "Create a central, searchable register of waste pickers organised by county.",
        "Issue a collision-safe identifier only after an administrator approves a profile.",
        "Give each approved picker a portable digital identity card in the Android application.",
        "Provide targeted broadcasts with recipient fan-out, inbox delivery and read tracking.",
        "Record material and weight so individual and community contribution becomes measurable.",
        "Provide live summaries and exportable operational reports for partner agencies.",
    ])
    doc.add_heading("1.2 Actors and scope", level=2)
    add_table(doc, ["Actor", "Primary responsibilities"], [
        ["Waste picker", "Self-register, sign in by phone, maintain profile, view identity, receive messages and record collections"],
        ["Administrator", "Review registrations, approve/reject/suspend, assign roles, broadcast messages, monitor activity and generate reports"],
        ["Supabase services", "Provide managed authentication, PostgreSQL persistence and profile photograph storage"],
    ])
    add_image(doc, GENERATED / "use_case.png", "Figure 5. Use cases supported by the implemented mobile and administrative interfaces.", width=6.6)
    doc.add_heading("1.3 Functional and non-functional requirements", level=2)
    add_table(doc, ["Area", "Implemented requirement"], [
        ["Identity", "Normalise Kenyan phone numbers, register once and issue a permanent ID after approval"],
        ["Administration", "Search/filter profiles, control status and role, and protect actions by administrator role"],
        ["Communication", "Broadcast to all approved pickers, one county or one individual; expose delivery and read totals"],
        ["Activity", "Capture material, positive weight and date; calculate personal and community totals"],
        ["Reporting", "Provide overview, registrations, regional, communication and collection summaries plus CSV/PDF outputs"],
        ["Security", "Use HTTPS, server-held service credentials, Supabase Auth, row-level security and server-side role guards"],
        ["Deployability", "Serve the React build and Express API from one Vercel origin with environment-managed configuration"],
    ])

    heading_page(doc, "2. System architecture and data flow")
    add_paragraph(doc, "The implementation follows a four-tier structure. The presentation layer contains the Android client and React administrative dashboard. The application layer is a Node.js/Express REST API containing validation, authorisation, identity workflow and communication fan-out. Supabase Auth and Storage form the integration layer, while PostgreSQL tables, views, triggers and storage buckets form the data layer.")
    add_image(doc, GENERATED / "architecture.png", "Figure 6. Four-tier architecture and the implemented components in each layer.", width=6.65)
    doc.add_heading("2.1 Single-origin deployment", level=2)
    add_paragraph(doc, "The dashboard and API share one HTTPS origin. Vercel serves the Vite build at the root and rewrites /api/* to a serverless function exporting the Express application. Browser requests therefore remain same-origin, and the Android client has one stable production base URL.")
    doc.add_heading("2.2 Information movement", level=2)
    add_image(doc, GENERATED / "data_flow.png", "Figure 7. Operational data flow without invented stores or payment functions.", width=6.65)
    add_table(doc, ["Request stage", "Responsibility"], [
        ["Client", "Collect input, attach the access token and render status-aware responses"],
        ["Express middleware", "Parse JSON, authenticate the token and require the appropriate actor/role"],
        ["Route handler", "Validate business input and orchestrate database, Auth and Storage operations"],
        ["Supabase", "Persist relational records, execute triggers/views and return authorised results"],
        ["Response", "Return JSON to the dashboard/mobile client; errors are converted to stable API messages"],
    ])
    doc.add_heading("2.3 Technology stack", level=2)
    add_table(doc, ["Layer", "Technology", "Purpose"], [
        ["Mobile", "Kotlin, XML, Retrofit/OkHttp, Coroutines", "Native waste picker workflow and REST client"],
        ["Web", "React 18, Vite, React Router, Recharts", "Administrative interface and analytics"],
        ["API", "Node.js, Express", "Business rules, authentication and reporting endpoints"],
        ["Platform", "Supabase Auth, PostgreSQL, Storage", "Identity, relational data, views, triggers and photographs"],
        ["Hosting", "Vercel", "Static dashboard and Express serverless function"],
        ["Version control", "Git and GitHub", "Source history and deployment hand-off"],
    ])

    heading_page(doc, "3. Database design and unique identity generation")
    add_image(doc, GENERATED / "erd.png", "Figure 8. Core tables and relationships implemented in supabase/schema.sql.", width=6.7)
    doc.add_heading("3.1 Core tables", level=2)
    add_table(doc, ["Table", "Purpose", "Important fields"], [
        ["regions", "Ten operating counties and their ID codes", "id, name, code"],
        ["waste_pickers", "Profile and approval lifecycle", "id, picker_id, phone, region_id, status, role"],
        ["admins", "Dashboard identities and privilege level", "id, email, role, is_active"],
        ["announcements", "Broadcast content and audience definition", "title, body, audience, region_id, recipient_count"],
        ["announcement_recipients", "Mobile inbox and read receipts", "announcement_id, picker_id, read_at"],
        ["collections", "Material collection records", "picker_id, material, weight_kg, collected_on"],
        ["device_tokens", "Optional push-notification targets", "picker_id, token, platform"],
        ["id_sequences", "Atomic counter per region and year", "region_code, year, last_value"],
    ])
    doc.add_heading("3.2 Reporting views", level=2)
    add_table(doc, ["View", "Output"], [
        ["v_registrations_by_region", "Total, approved, pending, rejected and suspended counts for every county"],
        ["v_communication_log", "Broadcast audience, sender, recipient count and read count"],
        ["v_collection_summary", "Per-picker record count and total kilograms, ranked by contribution"],
    ])
    doc.add_heading("3.3 Unique identity algorithm", level=2)
    add_paragraph(doc, "A registration begins without a picker ID. When its status first becomes approved, a PostgreSQL trigger calls generate_picker_id(region_id). The function locks and increments the id_sequences row for that region and calendar year inside the same transaction, then formats the result as WP-<REGION CODE>-<YEAR>-<FOUR DIGITS>. This prevents duplicate IDs during simultaneous approvals.")
    add_image(doc, GENERATED / "registration_flow.png", "Figure 9. Registration, approval and database-triggered identity assignment.", width=5.0)
    add_table(doc, ["Example", "Meaning"], [
        ["WP-KSM-2026-0001", "First approved Kisumu record in 2026"],
        ["WP-KKG-2026-0001", "First approved Kakamega record in 2026"],
        ["WP-SIA-2026-0003", "Third approved Siaya record in 2026"],
    ])

    heading_page(doc, "4. Authentication, authorisation and security model")
    doc.add_heading("4.1 Phone-first picker authentication", level=2)
    add_paragraph(doc, "Pickers sign in with a familiar Kenyan phone number. The API normalises supported forms to the 2547… or 2541… representation and deterministically maps that value to an internal email identity required by Supabase Auth. The internal address is never presented as the user's login name.")
    doc.add_heading("4.2 Administrator authentication", level=2)
    add_paragraph(doc, "Administrators authenticate with email and password through Supabase Auth. A successful Auth session is not sufficient by itself: the API also requires a matching active row in admins. Every dashboard route then applies both authenticate and requireAdmin middleware.")
    doc.add_heading("4.3 Role and state enforcement", level=2)
    add_table(doc, ["Control", "Enforcement"], [
        ["Access token", "Bearer token verified server-side for every protected request"],
        ["Actor type", "A picker cannot call administrative routes; an admin profile is required"],
        ["Admin role", "Superadmin-only actions are checked by middleware, not hidden only in the UI"],
        ["Picker status", "Pending, rejected and suspended profiles cannot submit collection activity"],
        ["Secrets", "Supabase service-role key remains in server/Vercel environment variables"],
        ["Database", "Row-level security prevents direct anonymous table access"],
    ])
    doc.add_heading("4.4 Communication delivery", level=2)
    add_paragraph(doc, "Creating an announcement inserts the message and resolves its audience to approved recipients. One announcement_recipients row is created for each target. These rows power the mobile inbox and unread badge; opening an item sets read_at, which feeds the communication report. Optional device-token delivery can supplement the in-app path when a push provider is configured.")

    heading_page(doc, "5. Administrative dashboard", "Verified against the live deployment; sensitive dashboard identifiers are masked")
    dashboard = [
        ("01-login.png", "Figure 10. Administrator sign-in."),
        ("02-dashboard.png", "Figure 11. Live dashboard KPIs and Recharts analytics."),
        ("03-waste-pickers.png", "Figure 12. Searchable and filterable waste picker register."),
        ("04-picker-detail.png", "Figure 13. Individual profile, status, role and activity detail."),
        ("05-communication.png", "Figure 14. Targeted broadcast composer and message history."),
        ("06-activity.png", "Figure 15. Community collection activity with totals and filters."),
        ("07-reports-registrations.png", "Figure 16. Registration report and filters."),
        ("08-reports-by-region.png", "Figure 17. County-level registration status report."),
        ("09-reports-communication.png", "Figure 18. Communication reach report."),
        ("10-reports-activity.png", "Figure 19. Per-picker community activity report."),
    ]
    add_paragraph(doc, "The dashboard is the coordinator's operational workspace. It exposes community KPIs, approval and role controls, communication tools, activity monitoring, CSV exports and the downloadable detailed PDF. The following evidence captures the implemented live screens.")
    for index, (filename, caption) in enumerate(dashboard):
        add_image(doc, SCREENSHOTS / filename, caption, width=6.65, page_break=index < len(dashboard) - 1)

    heading_page(doc, "6. Mobile application", "Physical-device evidence from the Android waste picker workflow")
    add_paragraph(doc, "The mobile application is designed for the waste picker rather than the coordinator. Session resolution directs the user to welcome, pending, rejected/suspended or approved content. Approved users receive a two-sided identity card, announcement inbox, personal activity record and editable profile.")
    mobile = [
        (("m01-splash.jpg", "Figure 20. Splash/session resolution."), ("m02-welcome.jpg", "Figure 21. Welcome and entry choices.")),
        (("m03-register.jpg", "Figure 22. Structured self-registration."), ("m04-signin.jpg", "Figure 23. Phone-number sign-in.")),
        (("m05-pending.jpg", "Figure 24. Pending approval status."), ("m06-id-front.jpg", "Figure 25. Digital ID card front.")),
        (("m07-id-back.jpg", "Figure 26. Digital ID card back."), ("m08-message.jpg", "Figure 27. Received announcement.")),
        (("m09-activity.jpg", "Figure 28. Personal collection history."), ("m10-profile.jpg", "Figure 29. Profile and photograph.")),
    ]
    for index, (left, right) in enumerate(mobile):
        add_mobile_pair(doc, (SCREENSHOTS / left[0], left[1]), (SCREENSHOTS / right[0], right[1]))
        if index < len(mobile) - 1:
            doc.add_page_break()
    doc.add_heading("6.1 Status-aware experience", level=2)
    add_bullets(doc, [
        "Pending users receive a clear explanation and a refresh action instead of inaccessible features.",
        "Approval reveals the database-issued picker ID and unlocks identity/activity functions.",
        "Rejected and suspended records receive explicit state messaging.",
        "Phone is kept read-only in profile editing because it is the account identifier.",
        "Profile photographs are stored in the configured Supabase Storage bucket and rendered on the identity card.",
    ])

    heading_page(doc, "7. Reporting and operational analytics")
    add_paragraph(doc, "The reporting module supports operational decision-making at two levels. The overview endpoint supplies live KPIs and graphs; detailed endpoints provide full registration, regional, communication and per-picker collection records. Client-side CSV exports support further analysis, while the generated PDF packages dated evidence, architecture and system summary for formal submission.")
    add_table(doc, ["Endpoint", "Purpose", "Completeness control"], [
        ["GET /api/reports/overview", "Counts, region totals, material totals and six-month trend", "Stable pagination for underlying collection/registration rows"],
        ["GET /api/reports/registrations", "Filterable picker register", "Validated filters and stable multi-page retrieval"],
        ["GET /api/reports/by-region", "County status totals", "Database aggregate view"],
        ["GET /api/reports/communication", "Broadcast delivery/read performance", "Stable multi-page retrieval from aggregate view"],
        ["GET /api/reports/collections", "Per-picker trips and kilograms", "Stable multi-page retrieval from aggregate view"],
    ])
    doc.add_heading("7.1 Management interpretation", level=2)
    top_region = max(data["by_region"], key=lambda row: row["total"])
    top_material = max(data["by_material"], key=lambda row: row["kg"])
    add_bullets(doc, [
        f"{top_region['region']} has the largest registered group ({top_region['total']} records); expansion can focus on counties with zero or one record.",
        f"{top_material['material']} is the largest material category ({top_material['kg']:,.2f} kg), useful for planning aggregation and buyers.",
        f"The pending queue contains {data['pending']} records and should be reviewed promptly so eligible pickers can receive identities.",
        f"Only {data['messages_read']} of {data['messages_delivered']} recorded deliveries has been opened; message timing, language and follow-up should be evaluated.",
        "The dated snapshot is evidence, not a permanent claim: the included capture script refreshes all source endpoints before a new report build.",
    ])

    heading_page(doc, "8. Deployment and operations")
    add_paragraph(doc, "The repository root is configured as one Vercel project. During deployment, the root install restores API dependencies and the vercel-build script builds the dashboard. Vercel serves dashboard/dist and routes every /api/* request to api/index.js, which exports the shared Express application.")
    add_table(doc, ["Deployment setting", "Configured value"], [
        ["Platform", "Vercel"],
        ["Build command", "npm run vercel-build"],
        ["Dashboard output", "dashboard/dist"],
        ["API entry", "api/index.js"],
        ["API rewrite", "/api/(.*) → /api/index"],
        ["Production origin", "https://waste-picker-system.vercel.app"],
        ["Database platform", "Supabase PostgreSQL"],
    ])
    doc.add_heading("8.1 Environment configuration", level=2)
    add_paragraph(doc, "Production values are stored in Vercel, never committed. The API requires SUPABASE_URL, SUPABASE_ANON_KEY and SUPABASE_SERVICE_ROLE_KEY, with PHONE_EMAIL_DOMAIN and PHOTO_BUCKET configuring identity mapping and image storage. Local demonstration credentials are explicit ignored environment values; no default password is embedded in source or this report.")
    doc.add_heading("8.2 Report publication", level=2)
    add_paragraph(doc, "The Vite build plugin copies the final PDF into the resolved build output at docs/waste-picker-system-report.pdf. The authenticated Reports page links to that artifact. It is intentionally public academic evidence and therefore excludes credentials; dashboard phone and national identity values are masked in the captured screens.")
    doc.add_heading("8.3 Rebuild sequence", level=2)
    add_numbered(doc, [
        "Set temporary WPS_REPORT_EMAIL and WPS_REPORT_PASSWORD values in the current shell.",
        "Run docs/capture_report_snapshot.ps1 to refresh all reporting endpoints into the JSON snapshot.",
        "Run python docs/generate_system_report.py to reproduce charts, diagrams and DOCX.",
        "Run python docs/export_report_pdf.py to export the final PDF through a dedicated Microsoft Word instance.",
        "Run the production dashboard build and verify the copied PDF hash and /api/health endpoint.",
    ])

    heading_page(doc, "9. Verification and testing")
    add_paragraph(doc, "The core workflow was exercised against the live deployment, and the completed report build adds repository-level structural checks. The table below summarises the end-to-end acceptance path.")
    add_table(doc, ["#", "Verification step", "Expected/verified result"], [
        [1, "Administrator signs in", "Active administrator session returned"],
        [2, "New picker registers", "Profile created as pending with no picker ID"],
        [3, "Pending picker attempts activity", "Request blocked by account status"],
        [4, "Administrator opens review queue", "Pending records listed"],
        [5, "Administrator approves", "Atomic WP-region-year-sequence ID assigned"],
        [6, "Approved picker records collection", "Positive material weight accepted"],
        [7, "Administrator broadcasts to a county", "Recipient rows created for approved targets"],
        [8, "Message reaches picker inbox", "Unread item visible"],
        [9, "Picker opens message", "read_at stored and unread count reduced"],
        [10, "Dashboard/report refresh", "Counts, kilograms and communication totals updated"],
        [11, "Picker calls administrator endpoint", "Administrator guard denies access"],
    ])
    doc.add_heading("9.1 Build and artifact checks", level=2)
    add_bullets(doc, [
        "Vite production build completes successfully and publishes the PDF into the actual resolved output directory.",
        "Modified API and seed JavaScript files pass Node syntax checking.",
        "Every Markdown image reference resolves to a committed evidence or generated file.",
        "The DOCX contains all nine generated visuals and twenty verified interface screenshots.",
        "The exported PDF is parsed for page count, required headings, final conclusion, EOF and absence of published seed passwords.",
        "The deployed health endpoint returns HTTP 200 before and after production publication.",
    ])

    heading_page(doc, "10. Security assessment")
    add_table(doc, ["Risk", "Implemented control", "Residual consideration"], [
        ["Unauthorised database access", "RLS on all tables; service key only on API", "Continue periodic policy review"],
        ["Privilege escalation", "Server-side actor and role middleware", "Add automated role-matrix tests"],
        ["Duplicate identity", "Atomic database sequence and trigger", "Monitor manual database changes"],
        ["Credential disclosure", "Environment-only secrets; published defaults removed", "Rotate any credential previously used in public history"],
        ["Sensitive screenshots", "Dashboard identifiers masked; report excludes credentials", "Use synthetic evidence for future publications"],
        ["Transport interception", "Single HTTPS production origin", "Maintain current TLS platform"],
        ["Oversized/incomplete reports", "Stable pagination and database aggregate views", "Move very large aggregates to dedicated SQL/RPC functions"],
        ["Profile photograph abuse", "Authenticated upload path and configured bucket", "Add content scanning and stricter file validation"],
    ])
    add_paragraph(doc, "Important operational action: any administrator password that was previously published or reused must be rotated. Removing a value from the current source does not remove it from Git history or invalidate an existing Supabase credential.")

    heading_page(doc, "11. Limitations and future work")
    add_table(doc, ["Limitation", "Recommended next step"], [
        ["Offline collection capture", "Queue records locally and synchronise idempotently when connectivity returns"],
        ["Push delivery while app is closed", "Complete Firebase/APNs provider configuration and delivery telemetry"],
        ["Low recorded message read rate", "Add language choice, delivery reminders and engagement follow-up"],
        ["National ID is not externally verified", "Integrate an authorised verification service with consent and privacy review"],
        ["English-only interface", "Add Kiswahili and Dholuo localisation with user testing"],
        ["Growing aggregate volume", "Move overview totals to database-side functions/materialised summaries as scale requires"],
        ["No automated CI suite", "Add API unit/integration tests, frontend tests and deployment smoke checks in GitHub Actions"],
        ["Mobile source maintained separately", "Link/version the Android repository and align release tags with this backend"],
    ])
    add_paragraph(doc, "Payment or M-Pesa disbursement is not part of the implemented system and is intentionally excluded from the architecture, data flow and current-system conclusions. It should only enter a future scope after requirements, regulatory, security and audit design.")

    heading_page(doc, "12. Conclusion")
    add_paragraph(doc, "The Waste Picker Management System fulfils the senior-project objective of turning an informal, difficult-to-reach workforce into a coordinated and measurable community. The implemented platform covers self-registration, administrative approval, collision-safe identity issue, targeted communication, collection activity and partner-oriented reporting through two role-appropriate interfaces.")
    add_paragraph(doc, f"The dated operational snapshot demonstrates working value rather than a design-only proposal: {data['total_pickers']} registered people, {data['approved']} approved identities, {data['total_kg']:,.2f} kg recorded and live reporting across counties, materials and communication. The report's graphs are generated directly from that snapshot, while its screenshots and diagrams map to code and schema present in the repository.")
    add_paragraph(doc, "The system is ready for continued controlled deployment, subject to immediate rotation of any credential previously published, stronger automated testing and the future-work priorities identified above. Its present foundation is technically coherent, reproducible and suitable for demonstration, academic assessment and incremental operational growth.")

    heading_page(doc, "Appendix A. API and repository reference")
    add_table(doc, ["Method and route", "Purpose"], [
        ["POST /api/auth/register", "Create picker Auth identity and pending profile"],
        ["POST /api/auth/login", "Picker phone/password sign-in"],
        ["POST /api/auth/admin/login", "Administrator sign-in"],
        ["GET/PATCH /api/me/*", "Picker profile, inbox and activity operations"],
        ["GET/PATCH/DELETE /api/pickers/*", "Administrative profile lifecycle and roles"],
        ["GET/POST /api/announcements/*", "Broadcast creation, history and detail"],
        ["GET /api/collections", "Community collection records"],
        ["GET /api/reports/*", "Overview and detailed reporting datasets"],
        ["GET /api/health", "Deployment health and timestamp"],
    ])
    doc.add_heading("Repository layout", level=2)
    add_table(doc, ["Path", "Contents"], [
        ["api/", "Vercel serverless entry points"],
        ["server/src/", "Express application, routes, middleware and Supabase helpers"],
        ["dashboard/src/", "React administrative interface"],
        ["supabase/schema.sql", "Tables, views, trigger, RLS and storage configuration"],
        ["docs/", "Canonical documentation, generator, snapshot and final report"],
        ["docs/screenshots/", "Ten dashboard and ten mobile evidence captures"],
        ["docs/generated/", "Four operational charts and five system diagrams"],
    ])

    heading_page(doc, "Appendix B. Report provenance")
    add_paragraph(doc, f"Snapshot timestamp: {data['captured_at']}")
    add_paragraph(doc, "Authenticated source endpoints:")
    add_bullets(doc, data.get("sources", []))
    add_paragraph(doc, "The snapshot stores aggregate counts only; no access token or password is written to disk. Labels such as e_waste are converted to presentation form (E-waste) without altering values. Region, material and trend totals are validated to reconcile with headline figures before report generation.")
    add_paragraph(doc, "Generated assets:")
    add_bullets(doc, [
        "status_summary.png — approval-state distribution",
        "registrations_by_region.png — stacked county status totals",
        "material_mix.png — cumulative kilograms by material",
        "registration_trend.png — six calendar months of registrations",
        "use_case.png, architecture.png, data_flow.png, registration_flow.png and erd.png — diagrams constrained to implemented code/schema",
    ])
    add_paragraph(doc, "Excluded proposal artifacts: any diagram depicting a persistent report store, notification-log database not present in the schema, or M-Pesa/Daraja disbursement is not evidence of the current system and is not included.")

    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    doc.save(OUTPUT)


def validate(data: dict) -> None:
    chart_files = [
        "status_summary.png",
        "registrations_by_region.png",
        "material_mix.png",
        "registration_trend.png",
        "use_case.png",
        "architecture.png",
        "data_flow.png",
        "registration_flow.png",
        "erd.png",
    ]
    missing = [name for name in chart_files if not (GENERATED / name).is_file()]
    if missing:
        raise RuntimeError(f"Generated assets missing: {', '.join(missing)}")
    if sum(row["total"] for row in data["by_region"]) != data["total_pickers"]:
        raise RuntimeError("Region totals do not reconcile with total_pickers")
    if not math.isclose(sum(float(row["kg"]) for row in data["by_material"]), float(data["total_kg"]), abs_tol=0.01):
        raise RuntimeError("Material totals do not reconcile with total_kg")
    if sum(row["registrations"] for row in data["registration_trend"]) != data["total_pickers"]:
        raise RuntimeError("Trend totals do not reconcile with total_pickers")
    report = Document(OUTPUT)
    headings = [paragraph.text for paragraph in report.paragraphs if paragraph.style and paragraph.style.name.startswith("Heading")]
    required = ["0. Executive system summary", "5. Administrative dashboard", "6. Mobile application", "12. Conclusion"]
    absent = [heading for heading in required if heading not in headings]
    if absent:
        raise RuntimeError(f"Report headings missing: {', '.join(absent)}")
    if len(report.inline_shapes) < 29:
        raise RuntimeError(f"Expected at least 29 embedded visuals, found {len(report.inline_shapes)}")


def main() -> None:
    GENERATED.mkdir(parents=True, exist_ok=True)
    data = load_snapshot()
    generate_charts(data)
    generate_diagrams()
    build_report(data)
    validate(data)
    report = Document(OUTPUT)
    print(f"Generated {len(list(GENERATED.glob('*.png')))} visuals")
    print(f"Generated {OUTPUT} ({len(report.paragraphs)} paragraphs, {len(report.tables)} tables, {len(report.inline_shapes)} images)")


if __name__ == "__main__":
    main()
